import Analysis_Finnhub as analysis
import helper_functions as hf
import csv
import math
import matplotlib.pyplot as plt
import finnhub
from docx import Document
from docx.shared import Inches

#===== INPUT PARAMETERS =====#
sector = 'h'
key1 = 'bucae7f48v6oa2u4ng20'
key2 = 'btts0j748v6ojt2hie60'

buffett_url = "/Users/ezzatsuhaime/Desktop/dcf_{}.csv".format(sector)
dcf_url = "/Users/ezzatsuhaime/Desktop/dcf_{}.docx".format(sector)


#===== FINNHUB CLIENT =====#
finnhub_client = finnhub.Client(api_key=key2)

#===== US COMPANIES =====#
us_tickers = "Tickers/{}.csv".format(sector) # file path name

#===== FILES =====#
document = Document()
document.add_heading('Monte-Carlo Discounted Cash Flow Model and Buffett Analysis', 0)

#===== ANALYSIS CLASS =====#
ana = analysis.Analysis_Finnhub(finnhub_client, "Data/{}_balance_annual.csv".format(sector), 
                                        "Data/{}_income_ttm.csv".format(sector), 
                                        "Data/{}_income_annual.csv".format(sector),
                                        "Data/{}_cashflows_ttm.csv".format(sector), 
                                        "Data/{}_cashflows_annual.csv".format(sector), 
                                        "Data/{}_basic_financials.csv".format(sector),
                                        "Data/{}_estimates.csv".format(sector))


#=============== ANALYSIS ===============#
with open(us_tickers, "r", newline = '') as ticker_file:
    ticker_reader = csv.reader(ticker_file, delimiter=' ', quotechar='|')
    
    for row in ticker_reader:   
        ticker = row[0] # ticker
        
        data_array = ana.buffett(ticker) # get Buffett analysis
        
        if(data_array is not None):
            
            if(data_array["p_agr_min"] > 0): # if min projected annual growth is positive
                document.add_heading("TICKER: {}".format(ticker), level=1)
                
                try:
                    image_path = "Images/{}.png".format(ticker)
                    image_path_2 = "Images/projection_{}.png".format(ticker)
                    
                    # Monte Carlo Analysis
                    plt.hist(ana.monte_carlo_DCF(ticker), bins = 30, density = True, color = "r")
                    plt.savefig(image_path)
                    document.add_picture(image_path, width=Inches(3.00))
                    plt.clf()

                    # Values for chart
                    length = data_array["length"]
                    eps_array = data_array['eps_array']
                    cagr_eps = data_array["cagr_eps"]
                    eps1 = data_array["eps1"]
                    projected_eps =  eps1 * math.pow(1+cagr_eps, 10)
                    coef = data_array["coef"][0]
                    intercept =  data_array["intercept"]
                    final_y = intercept + length * coef
                    
                    
                    # Drawing lines for chart
                    plt.scatter(data_array['count_array'], eps_array)
                    plt.plot([0, length], [intercept, final_y])
                    plt.plot([length, length+10], [final_y, projected_eps])

                    # Save chart
                    plt.savefig(image_path_2)
                    document.add_picture(image_path_2, width=Inches(3.0))
                    plt.clf()

                    # Write data describing chart
                    paragraph = document.add_paragraph("R-Value ROE: {}".format(data_array['r_value_roe']))
                    paragraph = document.add_paragraph("R-Value EPS: {}".format(data_array['r_value_eps']))
                    paragraph = document.add_paragraph("CAGR EPS: {}".format(data_array['cagr_eps']))
                    paragraph = document.add_paragraph("PE Min: {}".format(data_array['pe_min']))
                    paragraph = document.add_paragraph("PE Max: {}".format(data_array['pe_max']))
                    paragraph = document.add_paragraph("Projected AGR Min: {}".format(data_array['p_agr_min']))
                    paragraph = document.add_paragraph("Projected AGR Max: {}".format(data_array['p_agr_max']))

                    document.add_page_break()
                    #document.add_paragraph("\n\n")

                except Exception as e:
                    print(ticker, " : Error - ", e)
                    #print(ticker"Something happened with", ticker)
                    document.add_paragraph("Something wrong happened")
                    plt.clf()

document.save(dcf_url)
ticker_file.close()
